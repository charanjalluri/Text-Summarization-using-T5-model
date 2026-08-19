import io
import os
import tempfile

from app.core.exceptions import FileExtractionError
from app.core.logging_config import get_logger

logger = get_logger(__name__)

# Safe imports for library-specific dependencies that rely on native binaries
try:
    import fitz  # PyMuPDF
except Exception as e:
    fitz = None
    logger.warning(f"PyMuPDF (fitz) could not be loaded (likely missing native C DLLs): {e}")

try:
    import pdfplumber
except Exception as e:
    pdfplumber = None
    logger.warning(f"pdfplumber could not be loaded: {e}")

try:
    import speech_recognition as sr
except Exception as e:
    sr = None
    logger.warning(f"speech_recognition could not be loaded: {e}")

try:
    from pydub import AudioSegment
except Exception as e:
    AudioSegment = None
    logger.warning(f"pydub (AudioSegment) could not be loaded: {e}")

try:
    import pytesseract
except Exception as e:
    pytesseract = None
    logger.warning(f"pytesseract could not be loaded: {e}")

try:
    from docx import Document
except Exception as e:
    Document = None
    logger.warning(f"python-docx (Document) could not be loaded: {e}")


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Detects file type based on extension and extracts text contents."""
    if not filename:
        raise FileExtractionError("Filename is empty or invalid.")

    ext = filename.split(".")[-1].lower()
    if ext in {"txt", "text"}:
        return extract_txt(file_content)
    elif ext == "md":
        return extract_md(file_content)
    elif ext == "pdf":
        return extract_pdf(file_content)
    elif ext == "docx":
        return extract_docx(file_content)
    elif ext in {"wav", "mp3", "m4a", "flac"}:
        return extract_audio(file_content, ext)
    elif ext in {"png", "jpg", "jpeg", "tiff", "bmp"}:
        return extract_image(file_content)
    else:
        raise FileExtractionError(f"Unsupported file format: .{ext}")


def extract_txt(file_content: bytes) -> str:
    """Extracts raw text from bytes by trying UTF-8 and Latin-1 encodings."""
    try:
        return file_content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_content.decode("latin-1")
        except Exception as e:
            raise FileExtractionError(f"Failed to decode TXT file: {e}") from e


def extract_md(file_content: bytes) -> str:
    """Extracts raw text from Markdown bytes."""
    return extract_txt(file_content)


def extract_pdf(file_content: bytes) -> str:
    """Extracts text from PDF using PyMuPDF (fitz) with pdfplumber fallback."""
    text_parts = []

    # 1. Try PyMuPDF
    if fitz is not None:
        try:
            logger.info("Extracting PDF content using PyMuPDF (fitz)")
            with fitz.open(stream=file_content, filetype="pdf") as doc:
                for page in doc:
                    text_parts.append(page.get_text() or "")
            extracted_text = "\n".join(text_parts).strip()
            if len(extracted_text) >= 10:
                return extracted_text
            logger.info("PyMuPDF returned very little text. Attempting pdfplumber fallback.")
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}. Trying pdfplumber fallback.")

    # 2. Try pdfplumber
    if pdfplumber is not None:
        try:
            logger.info("Extracting PDF content using pdfplumber")
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts).strip()
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}", exc_info=True)
            raise FileExtractionError(f"Failed to extract text from PDF: {e}") from e

    raise FileExtractionError(
        "No PDF extraction libraries are available in this environment. "
        "Please check your system dependencies."
    )


def extract_docx(file_content: bytes) -> str:
    """Extracts paragraph and table cell text from DOCX bytes."""
    if Document is None:
        raise FileExtractionError(
            "python-docx is unavailable or failed to load. "
            "Cannot parse DOCX file formats."
        )

    try:
        logger.info("Extracting DOCX content using python-docx")
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables content
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    cleaned_row = []
                    for val in row_text:
                        if not cleaned_row or val != cleaned_row[-1]:
                            cleaned_row.append(val)
                    if cleaned_row:
                        text_parts.append(" | ".join(cleaned_row))

        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}", exc_info=True)
        raise FileExtractionError(f"Failed to extract text from DOCX: {e}") from e


def extract_audio(file_content: bytes, extension: str) -> str:
    """Transcribes audio using SpeechRecognition (Google Web API) and pydub."""
    if AudioSegment is None or sr is None:
        raise FileExtractionError(
            "Audio transcription libraries (pydub or speech_recognition) are unavailable or failed to load. "
            "Cannot parse audio files."
        )

    recognizer = sr.Recognizer()
    temp_wav_path = None
    try:
        logger.info(f"Extracting and transcribing audio format: {extension}")
        audio_stream = io.BytesIO(file_content)
        audio = AudioSegment.from_file(audio_stream, format=extension)

        wav_buffer = io.BytesIO()
        audio.export(wav_buffer, format="wav")
        wav_buffer.seek(0)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(wav_buffer.getvalue())
            temp_wav_path = tmp.name

        with sr.AudioFile(temp_wav_path) as source:
            audio_data = recognizer.record(source)

        logger.info("Sending transcription request to Google Speech API")
        text = recognizer.recognize_google(audio_data)
        return text.strip()
    except sr.UnknownValueError:
        logger.warning("Google Speech Recognition was unable to understand audio content.")
        raise FileExtractionError("Audio file read successfully, but Speech Recognition could not understand the speech.")
    except sr.RequestError as e:
        logger.error(f"Google Speech Recognition service connection error: {e}")
        raise FileExtractionError(f"Speech Recognition service error: {e}")
    except Exception as e:
        logger.error(f"Audio transcription error: {e}", exc_info=True)
        raise FileExtractionError(f"Failed to transcribe audio file: {e}")
    finally:
        if temp_wav_path and os.path.exists(temp_wav_path):
            try:
                os.unlink(temp_wav_path)
                logger.info(f"Cleaned up temporary audio WAV file: {temp_wav_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up temporary audio WAV file {temp_wav_path}: {e}")


def extract_image(file_content: bytes) -> str:
    """Performs OCR text extraction from image bytes using pytesseract."""
    if pytesseract is None:
        raise FileExtractionError(
            "pytesseract library is unavailable or failed to load. "
            "Cannot perform image text extraction."
        )

    try:
        from PIL import Image
        logger.info("Extracting text from image using Tesseract OCR")
        image = Image.open(io.BytesIO(file_content))
        text = pytesseract.image_to_string(image)
        return text.strip()
    except pytesseract.TesseractNotFoundError:
        logger.error("Tesseract binary not found on local path.")
        raise FileExtractionError("Tesseract OCR binary was not found. Please install Tesseract on your system to process images.")
    except Exception as e:
        logger.error(f"Image OCR error: {e}", exc_info=True)
        raise FileExtractionError(f"Failed to extract text from image: {e}")
